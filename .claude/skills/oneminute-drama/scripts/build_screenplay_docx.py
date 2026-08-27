#!/usr/bin/env python3
"""ONE MINUTE STUDIO 하우스 서식으로 Word 문서를 만든다.

    python3 build_screenplay_docx.py spec.json out.docx

매번 docx 빌더를 새로 짜면 서식이 조금씩 달라져서 문서가 시리즈처럼 보이지 않는다.
서식은 여기 한 곳에만 두고, 내용만 JSON으로 넘긴다.

spec.json 구조 — cover 는 생략 가능, blocks 는 위에서부터 순서대로 렌더된다.

{
  "cover": {"kicker": "DIALOGUE SCRIPT", "title": "11:11",
            "subtitle": "ELEVEN ELEVEN", "sub2": "EP 01–10 · 대사 대본",
            "tagline": "매일 밤 11시 11분, 딱 1분.",
            "note": "그녀가 원할수록 그는 사라진다.",
            "sign": "ONE MINUTE STUDIO · @oneminute.studio"},
  "footer": "ONE MINUTE STUDIO · 11:11",
  "blocks": [
    {"type": "h1", "text": "1. 기획"},
    {"type": "h2", "text": "EP 01. 줄"},
    {"type": "h3", "text": "훅 카드"},
    {"type": "p",  "text": "본문", "bold": false, "color": "MUTED", "size": 18},
    {"type": "rule"},
    {"type": "pagebreak"},
    {"type": "table", "cols": [1400, 7600], "header": true, "zebra": true,
     "boldCol0": true, "size": 17,
     "rows": [["항목", "값"], ["시계", "11:11"]]},
    {"type": "scene", "no": "S#01", "place": "골목 / 밤", "tc": "0:00–0:03",
     "action": ["어둠 속에 수십 명이 줄지어 서 있다."],
     "se": ["바람. 그 아래로 먼 초침 소리."],
     "lines": [{"who": "도하", "kr": "안 켜집니다.", "en": "It won't record.",
                "dir": "★ 첫 대사"}],
     "sub": {"kr": "소원은 이뤄진다", "en": "WISHES COME TRUE"}}
  ]
}

색 키워드: INK(본문) MUTED(흐림) KEY(강조) RED(액센트 #E2342B). 직접 6자리 hex 도 된다.
"""
import json
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips

FONT = "Malgun Gothic"
W = 9000  # 본문 폭 (A4, 1인치 여백)
COLORS = {"INK": "1A1A1C", "MUTED": "6A6E72", "KEY": "A9560A", "RED": "E2342B"}
CENTER = WD_ALIGN_PARAGRAPH.CENTER


def _col(name):
    if not name:
        return COLORS["INK"]
    return COLORS.get(str(name).upper(), str(name).lstrip("#"))


def _shade(cell, fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(el)


def _border(p, edge="top", color="C7CED4", size=6):
    pPr = p._p.get_or_add_pPr()
    bd = pPr.find(qn("w:pBdr"))
    if bd is None:
        bd = OxmlElement("w:pBdr")
        pPr.append(bd)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(size))
    e.set(qn("w:color"), color)
    bd.append(e)


def para(doc, text="", size=20, bold=False, italic=False, color="INK",
         align=None, before=60, after=60, line=300):
    """size 는 half-point (docx 관례). 20 이면 10pt."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before / 20)
    pf.space_after = Pt(after / 20)
    if line:
        pf.line_spacing = Pt(line / 20)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(size / 2)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = None
    from docx.shared import RGBColor
    r.font.color.rgb = RGBColor.from_string(_col(color))
    return p


def heading(doc, text, level):
    size, color, before, after = {
        1: (30, "INK", 360, 140),
        2: (24, "INK", 280, 110),
        3: (21, "KEY", 200, 80),
    }[level]
    p = para(doc, text, size=size, bold=True, color=color, before=before, after=after)
    p.style = doc.styles[f"Heading {level}"]
    for r in p.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.size = Pt(size / 2)
        r.font.bold = True
        from docx.shared import RGBColor
        r.font.color.rgb = RGBColor.from_string(_col(color))
    return p


def table(doc, cols, rows, header=True, zebra=False, boldCol0=False, size=18):
    t = doc.add_table(rows=0, cols=len(cols))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    try:
        t.style = doc.styles["Table Grid"]
    except KeyError:
        pass
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].width = Twips(cols[j])
            fill = None
            if header and i == 0:
                fill = "E8EAEC"
            elif zebra and i % 2 == 0:
                fill = "F6F7F8"
            if fill:
                _shade(cells[j], fill)
            cp = cells[j].paragraphs[0]
            cp.paragraph_format.space_before = Pt(1)
            cp.paragraph_format.space_after = Pt(1)
            r = cp.add_run("" if val is None else str(val))
            r.font.name = FONT
            r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            r.font.size = Pt(size / 2)
            r.font.bold = (header and i == 0) or (boldCol0 and j == 0)
            from docx.shared import RGBColor
            r.font.color.rgb = RGBColor.from_string(
                _col("3B454F" if (header and i == 0) else "INK"))
    return t


def scene(doc, b):
    head = f"{b.get('no', '')}. {b.get('place', '')}".strip(". ")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(17)
    p.paragraph_format.space_after = Pt(6)
    from docx.shared import RGBColor
    for txt, sz, bold, col in ((head, 22, True, "INK"),
                               (f"      ({b['tc']})" if b.get("tc") else "", 19, True, "KEY")):
        if not txt:
            continue
        r = p.add_run(txt)
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.size = Pt(sz / 2)
        r.font.bold = bold
        r.font.color.rgb = RGBColor.from_string(_col(col))
    for a in b.get("action", []):
        para(doc, a, size=20, before=40, after=40)
    for s in b.get("se", []):
        para(doc, f"S.E.  {s}", size=18, italic=True, color="MUTED", before=60, after=60)
    for ln in b.get("lines", []):
        para(doc, ln.get("who", ""), size=19, bold=True, align=CENTER, before=140, after=20)
        para(doc, ln.get("kr", ""), size=22, align=CENTER, after=20)
        if ln.get("en"):
            para(doc, ln["en"], size=17, italic=True, color="MUTED", align=CENTER, after=20)
        if ln.get("dir"):
            para(doc, ln["dir"], size=16, color="MUTED", align=CENTER, after=60)
    sub = b.get("sub")
    if sub:
        para(doc, f"자막:  {sub.get('kr','')}", size=21, bold=True, color="RED",
             before=140, after=20)
        if sub.get("en"):
            para(doc, f"SUB:  {sub['en']}", size=17, color="MUTED", after=40)


def build(spec, out):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = Pt(10)

    cv = spec.get("cover")
    if cv:
        if cv.get("kicker"):
            para(doc, cv["kicker"], size=18, bold=True, color="MUTED",
                 align=CENTER, before=1600, after=0)
        para(doc, cv.get("title", ""), size=64, bold=True, align=CENTER, before=120, after=0)
        if cv.get("subtitle"):
            para(doc, cv["subtitle"], size=20, color="MUTED", align=CENTER, before=40, after=0)
        if cv.get("sub2"):
            para(doc, cv["sub2"], size=28, bold=True, align=CENTER, before=300, after=0)
        _border(para(doc, "", size=10, align=CENTER, before=420, after=0))
        if cv.get("tagline"):
            para(doc, cv["tagline"], size=24, bold=True, align=CENTER, before=280)
        if cv.get("note"):
            para(doc, cv["note"], size=20, color="MUTED", align=CENTER, before=40)
        if cv.get("sign"):
            para(doc, cv["sign"], size=18, bold=True, color="KEY", align=CENTER, before=520)
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    for b in spec.get("blocks", []):
        t = b.get("type", "p")
        if t == "pagebreak":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        elif t == "rule":
            _border(para(doc, "", size=10, before=b.get("before", 400), after=0))
        elif t in ("h1", "h2", "h3"):
            heading(doc, b["text"], int(t[1]))
        elif t == "table":
            table(doc, b.get("cols") or [W // max(1, len(b["rows"][0]))] * len(b["rows"][0]),
                  b["rows"], header=b.get("header", True), zebra=b.get("zebra", False),
                  boldCol0=b.get("boldCol0", False), size=b.get("size", 18))
        elif t == "scene":
            scene(doc, b)
        else:
            para(doc, b.get("text", ""), size=b.get("size", 20), bold=b.get("bold", False),
                 italic=b.get("italic", False), color=b.get("color", "INK"),
                 align=CENTER if b.get("center") else None,
                 before=b.get("before", 60), after=b.get("after", 60))

    if spec.get("footer"):
        _border(para(doc, "", size=10, before=600, after=0))
        para(doc, spec["footer"], size=16, color="MUTED", align=CENTER, before=160)

    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Pt(72)
    doc.save(out)
    return out


if __name__ == "__main__":
    if len(sys.argv) != 3:
        sys.exit("usage: build_screenplay_docx.py spec.json out.docx")
    with open(sys.argv[1], encoding="utf-8") as f:
        spec = json.load(f)
    print("written", build(spec, sys.argv[2]))
